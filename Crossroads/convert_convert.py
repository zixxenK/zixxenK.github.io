"""
Crossroads Proposal: Markdown -> Professional Word Document Converter
Produces a submission-ready .docx with:
  - Professional cover page
  - Table of contents placeholder
  - Properly numbered headings (H1=Title, H2=major sections, H3=subsections)
  - Formatted tables with shaded header rows
  - Inline bold, italic, and hyperlink rendering
  - Page breaks at horizontal rule (---) boundaries
  - Page numbers in footer
  - Times New Roman 12pt body, proper margins, 1.15 line spacing
"""

import re
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def force_font_name(element, name):
    """Force font name on an rPr or style element, clearing theme overrides."""
    rPr = element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        element.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    # Set all four font slots
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)
    # Remove theme attributes that override explicit font names
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        key = qn(attr)
        if rFonts.get(key) is not None:
            del rFonts.attrib[key]


def set_run_font(run, name="Times New Roman", size=12, bold=None, italic=None, color=None):
    """Utility to configure a run's font properties."""
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # Force the XML-level font to beat theme overrides
    force_font_name(run._r, name)


def set_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.15):
    """Set spacing on a paragraph."""
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def add_page_number(section):
    """Insert a centered page number into the section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # PAGE field code
    run1 = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1._r.append(fld_char_begin)
    run2 = p.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr)
    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char_end)
    for r in [run1, run2, run3]:
        set_run_font(r, size=10)


# ---------------------------------------------------------------------------
# Core conversion
# ---------------------------------------------------------------------------

def markdown_to_docx(md_path, docx_path):
    # Read source markdown
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    # Convert to HTML via the markdown library
    html = markdown.markdown(md_text, extensions=["tables", "smarty"])
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    # -- Global styles -------------------------------------------------------
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    # Force XML-level font on the Normal style so theme doesn't override
    force_font_name(style_normal.element, "Times New Roman")

    # Heading styles
    for level in range(1, 5):
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = "Times New Roman"
        hstyle.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            hstyle.font.size = Pt(18)
        elif level == 2:
            hstyle.font.size = Pt(14)
        elif level == 3:
            hstyle.font.size = Pt(12)
        else:
            hstyle.font.size = Pt(12)
        hstyle.font.bold = True
        # Strip theme font overrides from heading styles too
        force_font_name(hstyle.element, "Times New Roman")

    # -- Margins -------------------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # -----------------------------------------------------------------------
    # Build document from parsed HTML elements
    # -----------------------------------------------------------------------
    page_phase = "memo"  # memo -> title -> content

    def add_inline_content(paragraph, element):
        """Recursively add inline children (text, <strong>, <em>, <a>, etc.) to a paragraph."""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text:
                    run = paragraph.add_run(text)
                    set_run_font(run)
            elif isinstance(child, Tag):
                if child.name == "br":
                    run = paragraph.add_run()
                    run.add_break()
                elif child.name in ("strong", "b"):
                    # Could contain nested <em>
                    for sub in child.children:
                        if isinstance(sub, NavigableString):
                            run = paragraph.add_run(str(sub))
                            set_run_font(run, bold=True)
                        elif isinstance(sub, Tag) and sub.name in ("em", "i"):
                            run = paragraph.add_run(sub.get_text())
                            set_run_font(run, bold=True, italic=True)
                        else:
                            run = paragraph.add_run(sub.get_text())
                            set_run_font(run, bold=True)
                elif child.name in ("em", "i"):
                    for sub in child.children:
                        if isinstance(sub, NavigableString):
                            run = paragraph.add_run(str(sub))
                            set_run_font(run, italic=True)
                        elif isinstance(sub, Tag) and sub.name in ("strong", "b"):
                            run = paragraph.add_run(sub.get_text())
                            set_run_font(run, bold=True, italic=True)
                        else:
                            run = paragraph.add_run(sub.get_text())
                            set_run_font(run, italic=True)
                elif child.name == "a":
                    link_text = child.get_text()
                    href = child.get("href", "")
                    run = paragraph.add_run(link_text)
                    set_run_font(run, color=RGBColor(0x05, 0x63, 0xC1))
                    run.underline = True
                    # Embed actual hyperlink via oxml
                    try:
                        r_id = paragraph.part.relate_to(
                            href,
                            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                            is_external=True,
                        )
                        hyperlink = parse_xml(
                            f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                            f"</w:hyperlink>"
                        )
                        # Move the run element into the hyperlink
                        hyperlink.append(run._r)
                        paragraph._p.append(hyperlink)
                    except Exception:
                        pass  # Fallback: the underlined blue text is still present
                else:
                    # Generic inline tag — just grab text
                    run = paragraph.add_run(child.get_text())
                    set_run_font(run)

    def process_element(el):
        nonlocal page_phase

        if not isinstance(el, Tag):
            return

        # --- Horizontal rules -> page breaks ---
        if el.name == "hr":
            doc.add_page_break()
            if page_phase == "memo":
                page_phase = "title"
            elif page_phase == "title":
                page_phase = "content"
            return

        # --- Headings ---
        if el.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(el.name[1])
            text = el.get_text()

            # Memo page: center the MEMORANDUM / LETTER OF TRANSMITTAL heading
            if page_phase == "memo" and ("memorandum" in text.lower() or "letter of transmittal" in text.lower()):
                for _ in range(2):
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.space_before = Pt(0)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text.upper())
                set_run_font(run, size=16, bold=True)
                set_paragraph_spacing(p, before=12, after=12)
                return

            if level == 1 and page_phase == "title":
                # === BUILD TITLE PAGE ===
                # Add vertical space to centre the title block
                for _ in range(6):
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.space_before = Pt(0)

                title_p = doc.add_paragraph()
                title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = title_p.add_run(text.upper())
                set_run_font(run, size=22, bold=True)
                set_paragraph_spacing(title_p, after=24)
                return

            # Title page: center other headings (subtitle)
            if page_phase == "title":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                set_run_font(run, size=14 if level <= 2 else 12, bold=True)
                set_paragraph_spacing(p, before=6, after=6)
                return

            heading = doc.add_heading(text, level=min(level, 4))
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Force font on heading runs
            for run in heading.runs:
                set_run_font(run, size={1: 18, 2: 14, 3: 12}.get(level, 12), bold=True)
            return

        # --- Paragraphs ---
        if el.name == "p":
            text = el.get_text()
            if page_phase == "memo":
                # Memo page: left-aligned with inline formatting
                p = doc.add_paragraph()
                add_inline_content(p, el)
                set_paragraph_spacing(p, before=2, after=2)
                return

            if page_phase == "title":
                # Title page: centered
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_content(p, el)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                set_paragraph_spacing(p, before=4, after=4)
                return

            p = doc.add_paragraph()
            add_inline_content(p, el)
            set_paragraph_spacing(p, before=0, after=6)
            return

        # --- Unordered / Ordered lists ---
        if el.name in ("ul", "ol"):
            list_style = "List Bullet" if el.name == "ul" else "List Number"
            for li in el.find_all("li", recursive=False):
                p = doc.add_paragraph(style=list_style)
                add_inline_content(p, li)
                set_paragraph_spacing(p, before=0, after=3)
            return

        # --- Tables ---
        if el.name == "table":
            rows = el.find_all("tr")
            if not rows:
                return
            col_count = max(
                len(r.find_all(["th", "td"])) for r in rows
            )
            tbl = doc.add_table(rows=0, cols=col_count)
            tbl.style = "Table Grid"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            for row_idx, row_el in enumerate(rows):
                cells_el = row_el.find_all(["th", "td"])
                row_obj = tbl.add_row()
                for col_idx, cell_el in enumerate(cells_el):
                    if col_idx >= col_count:
                        break
                    cell = row_obj.cells[col_idx]
                    cell.text = ""  # clear default
                    p = cell.paragraphs[0]
                    # Render inline content inside cells
                    add_inline_content(p, cell_el)
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.name = "Times New Roman"

                    # Header row styling
                    is_header = row_idx == 0 or cell_el.name == "th"
                    if is_header:
                        set_cell_shading(cell, "2E4057")  # dark blue-grey
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.bold = True

            # Small spacing after table
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)
            return

    # --- Walk top-level elements ---
    for child in soup.children:
        process_element(child)

    # -- Page numbers on every section ---
    for section in doc.sections:
        add_page_number(section)

    # -- Save ---------------------------------------------------------------
    doc.save(docx_path)
    print(f"[OK] Saved {docx_path}")


# ---------------------------------------------------------------------------
# Run
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    md_file = r"c:\Documents\Crossroads\workinprogressenglishproposalcrossroads.md"
    docx_file = r"c:\Documents\Crossroads\Crossroads_Proposal_v6.docx"
    markdown_to_docx(md_file, docx_file)
